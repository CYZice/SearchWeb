import io
import json
import os
import shutil
import sys
import tempfile
import unittest

from docx import Document
from docx.shared import Inches
from PIL import Image
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if ROOT_DIR not in sys.path:
    sys.path.insert(0, ROOT_DIR)

import app.main as app_main
from app.models import Base, Inscription
from scripts import word_parser


def field_name(name):
    for key, value in word_parser.FIELD_MAP.items():
        if value == name:
            return key
    raise AssertionError(f"Missing field mapping for {name}")


def image_bytes(size, color, fmt="JPEG"):
    img = Image.new("RGB", size, color=color)
    buffer = io.BytesIO()
    img.save(buffer, format=fmt)
    buffer.seek(0)
    return buffer


def add_record(doc, serial, name, transcript, include_small=False, include_large=False):
    doc.add_paragraph(f"[{serial}]")
    doc.add_paragraph(f"{field_name('name')}:{name}")
    doc.add_paragraph(f"{field_name('transcript')}:{transcript}")
    if include_small:
        doc.add_paragraph().add_run().add_picture(
            image_bytes((32, 32), "red"), width=Inches(0.2)
        )
    if include_large:
        doc.add_paragraph().add_run().add_picture(
            image_bytes((1800, 1800), "blue"), width=Inches(1.5)
        )


class WordImportStreamingTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.mkdtemp(prefix="word_import_test_")
        self.images_dir = os.path.join(self.temp_dir, "images")
        self.chars_dir = os.path.join(self.images_dir, "chars")
        self.conflicts_dir = os.path.join(self.temp_dir, "conflicts")
        os.makedirs(self.chars_dir, exist_ok=True)

        self.old_paths = {
            "IMAGES_DIR": word_parser.IMAGES_DIR,
            "CHAR_IMAGES_DIR": word_parser.CHAR_IMAGES_DIR,
            "CONFLICTS_DIR": word_parser.CONFLICTS_DIR,
        }
        word_parser.IMAGES_DIR = self.images_dir
        word_parser.CHAR_IMAGES_DIR = self.chars_dir
        word_parser.CONFLICTS_DIR = self.conflicts_dir

        self.engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(bind=self.engine)
        Session = sessionmaker(bind=self.engine)
        self.db = Session()

    def tearDown(self):
        self.db.close()
        for name, value in self.old_paths.items():
            setattr(word_parser, name, value)
        shutil.rmtree(self.temp_dir)

    def make_docx(self, records):
        doc = Document()
        for args in records:
            add_record(doc, **args)
        path = os.path.join(self.temp_dir, "sample.docx")
        doc.save(path)
        return path

    def test_streaming_import_saves_records_and_images(self):
        path = self.make_docx(
            [
                {
                    "serial": "001",
                    "name": "Record One",
                    "transcript": "Transcript One",
                    "include_small": True,
                    "include_large": True,
                },
                {
                    "serial": "002",
                    "name": "Record Two",
                    "transcript": "Transcript Two",
                },
            ]
        )

        records_iter = word_parser.iter_docx_records(path)
        self.assertFalse(isinstance(records_iter, list))
        self.assertEqual(len(word_parser.parse_docx(path)), 2)

        result = word_parser.process_import(path, self.db, batch_size=1)
        self.assertEqual(result["success"], 2)
        self.assertEqual(result["skipped"], 0)
        self.assertEqual(self.db.query(Inscription).count(), 2)

        first = (
            self.db.query(Inscription)
            .filter(Inscription.serial_num == "[001]")
            .one()
        )
        image_urls = json.loads(first.image_url)
        self.assertEqual(len(image_urls), 1)
        self.assertTrue(image_urls[0].startswith(word_parser.IMAGE_URL_PREFIX))
        self.assertTrue(os.listdir(self.chars_dir))
        self.assertTrue(
            os.path.exists(
                os.path.join(self.images_dir, os.path.basename(image_urls[0]))
            )
        )

    def test_duplicate_serials_are_skipped_by_serial_number(self):
        path = self.make_docx(
            [
                {"serial": "001", "name": "Record One", "transcript": "A"},
                {"serial": "001", "name": "Record One Duplicate", "transcript": "B"},
            ]
        )

        result = word_parser.process_import(path, self.db, batch_size=2)
        self.assertEqual(result["success"], 1)
        self.assertEqual(result["skipped"], 1)
        self.assertEqual(result["skipped_items"][0]["reason"], "Duplicate Serial Number (in same document)")
        self.assertEqual(self.db.query(Inscription).count(), 1)

    def test_database_duplicate_gets_conflict_token_and_can_overwrite(self):
        existing = Inscription(serial_num="[001]", name="Old", transcript="Old Text")
        self.db.add(existing)
        self.db.commit()

        path = self.make_docx(
            [{"serial": "001", "name": "New", "transcript": "New Text"}]
        )
        result = word_parser.process_import(path, self.db)

        self.assertEqual(result["success"], 0)
        self.assertEqual(result["skipped"], 1)
        conflict = result["skipped_items"][0]
        self.assertEqual(conflict["existing_id"], existing.id)
        self.assertIn("conflict_token", conflict)
        self.assertNotIn("new_data", conflict)

        overwrite = app_main.overwrite_inscription(
            {"existing_id": existing.id, "conflict_token": conflict["conflict_token"]},
            db=self.db,
        )
        self.assertEqual(overwrite["status"], "success")
        updated = self.db.query(Inscription).filter(Inscription.id == existing.id).one()
        self.assertEqual(updated.name, "New")
        self.assertEqual(updated.transcript, "New Text")
        self.assertFalse(
            os.path.exists(word_parser.conflict_path(conflict["conflict_token"]))
        )


if __name__ == "__main__":
    unittest.main()
