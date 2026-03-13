import os
import sys
from docx import Document
from docx.shared import Inches

# Ensure data/raw_word exists
os.makedirs("data/raw_word", exist_ok=True)

doc = Document()

# Record 1: 【001】
doc.add_paragraph("【001】")
doc.add_paragraph("器名：耶律羽之墓志")
doc.add_paragraph("时代：会同五年（公元 942 年）")
doc.add_paragraph("别称：无")
doc.add_paragraph("出土、发现：1992 年 7 月赤峰市阿鲁科尔沁旗罕苏木苏木朝克图山")
doc.add_paragraph("现藏：内蒙古文物考古研究所")
doc.add_paragraph("主要著录：《辽代石刻文续编》2010:3-6；《内蒙古辽代石刻文研究》2002：2-16；《文物》1996（1）：4-32；")
doc.add_paragraph("形制：石质为灰色砂岩，正面与四边素面磨光。楷书。38 行，共 1210 字。拓片长 112 厘米、宽 103 厘米，厚 12 厘米")
doc.add_paragraph("释文：")
doc.add_paragraph("大契丹国东京太傅相公墓誌铭并序。")
doc.add_paragraph("蓟门那明远撰并书。")
doc.add_paragraph("夫欲建皇极庶垂风，必资栋梁之材，更赖盐梅之士。其或非熊应兆，卧龙见称。时推命世之智，代许间生之杰。")

# Add a dummy image (we'll just create a small red square)
from PIL import Image
import io

img = Image.new('RGB', (100, 100), color = 'red')
img_byte_arr = io.BytesIO()
img.save(img_byte_arr, format='JPEG')
img_byte_arr.seek(0)

doc.add_paragraph().add_run().add_picture(img_byte_arr, width=Inches(1.25))

# Record 2: 【002】
doc.add_paragraph("【002】")
doc.add_paragraph("器名：刘存规墓志")
doc.add_paragraph("时代：应历五年（公元 955 年）")
doc.add_paragraph("出土：不详")
doc.add_paragraph("释文：")
doc.add_paragraph("刘公墓志铭。")

# Add another image (blue)
img2 = Image.new('RGB', (100, 100), color = 'blue')
img2_byte_arr = io.BytesIO()
img2.save(img2_byte_arr, format='JPEG')
img2_byte_arr.seek(0)

doc.add_paragraph().add_run().add_picture(img2_byte_arr, width=Inches(1.25))


doc.save("data/raw_word/test_sample.docx")
print("Generated data/raw_word/test_sample.docx")
